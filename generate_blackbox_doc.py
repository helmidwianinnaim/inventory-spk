from docx import Document
from docx.shared import Inches

doc = Document()
doc.add_heading('Tabel Pengujian Blackbox Sistem SPK', 0)

def add_table(title, rows):
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=len(rows)+1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'No'
    hdr_cells[1].text = 'Skenario'
    hdr_cells[2].text = 'Hasil yang Diharapkan'
    hdr_cells[3].text = 'Status'
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.cell(i+1, j).text = str(val)
    doc.add_paragraph()

# 1. Login & Hak Akses
add_table('Login & Hak Akses', [
    ['1', 'Pengguna memasukkan username/email dan password yang salah', 'Sistem menampilkan pesan error "Cek username/email dan password anda"', 'Berhasil'],
    ['2', 'Pengguna memasukkan username/email dan password yang benar', 'Pengguna berhasil masuk ke dashboard', 'Berhasil'],
    ['3', 'Pengguna mengakses menu tanpa login', 'Sistem mengarahkan ke halaman login', 'Berhasil'],
    ['4', 'User biasa mencoba mengakses menu admin', 'Sistem menolak akses dan menampilkan pesan "Anda tidak memiliki akses"', 'Berhasil'],
    ['5', 'Pengguna mengubah password di halaman ganti password', 'Password berhasil diubah dan sistem menampilkan notifikasi', 'Berhasil'],
    ['6', 'Pengguna mengubah data profil', 'Data profil berhasil diubah dan notifikasi muncul', 'Berhasil'],
    ['7', 'Pengguna mencoba login dengan field kosong', 'Sistem menampilkan pesan error "Field wajib diisi"', 'Berhasil']
])

# 2. Master Data: Produk
add_table('Master Data: Produk', [
    ['1', 'Menampilkan data produk', 'Data produk berhasil tampil', 'Berhasil'],
    ['2', 'Menambah data produk', 'Data produk berhasil disimpan', 'Berhasil'],
    ['3', 'Menambah produk tanpa mengisi nama produk', 'Sistem menampilkan pesan error "Nama produk wajib diisi"', 'Berhasil'],
    ['4', 'Menambah produk dengan kode yang sudah ada', 'Sistem menolak penyimpanan dan menampilkan pesan error "Kode produk sudah digunakan"', 'Berhasil'],
    ['5', 'Mengedit data produk', 'Data produk berhasil diubah', 'Berhasil'],
    ['6', 'Mengedit produk dengan format harga tidak valid', 'Sistem menampilkan pesan error "Format harga tidak valid"', 'Berhasil'],
    ['7', 'Menghapus produk yang tidak berelasi', 'Data produk berhasil dihapus', 'Berhasil'],
    ['8', 'Menghapus produk yang masih berelasi dengan transaksi produk keluar', 'Sistem menolak penghapusan dan menampilkan pesan error "Data tidak dapat dihapus karena masih digunakan di transaksi lain"', 'Berhasil'],
    ['9', 'Melihat detail produk', 'Data detail produk berhasil tampil', 'Berhasil']
])

# 3. Master Data: Kategori
add_table('Master Data: Kategori', [
    ['1', 'Menampilkan data kategori', 'Data kategori berhasil tampil', 'Berhasil'],
    ['2', 'Menambah data kategori', 'Data kategori berhasil disimpan', 'Berhasil'],
    ['3', 'Menambah kategori tanpa nama', 'Sistem menampilkan pesan error "Nama kategori wajib diisi"', 'Berhasil'],
    ['4', 'Menambah kategori dengan nama yang sudah ada', 'Sistem menolak penyimpanan dan menampilkan pesan error "Kategori sudah ada"', 'Berhasil'],
    ['5', 'Mengedit data kategori', 'Data kategori berhasil diubah', 'Berhasil'],
    ['6', 'Menghapus data kategori', 'Data kategori berhasil dihapus', 'Berhasil']
])

# 4. Master Data: Bahan Baku
add_table('Master Data: Bahan Baku', [
    ['1', 'Menampilkan data bahan baku', 'Data bahan baku berhasil tampil', 'Berhasil'],
    ['2', 'Menambah data bahan baku', 'Data bahan baku berhasil disimpan', 'Berhasil'],
    ['3', 'Menambah bahan baku tanpa nama', 'Sistem menampilkan pesan error "Nama bahan baku wajib diisi"', 'Berhasil'],
    ['4', 'Mengedit data bahan baku', 'Data bahan baku berhasil diubah', 'Berhasil'],
    ['5', 'Menghapus bahan baku yang tidak berelasi', 'Data bahan baku berhasil dihapus', 'Berhasil'],
    ['6', 'Menghapus bahan baku yang masih berelasi dengan transaksi', 'Sistem menolak penghapusan dan menampilkan pesan error "Data tidak dapat dihapus karena masih digunakan di transaksi lain"', 'Berhasil']
])

# 5. Master Data: Formula Produk
add_table('Master Data: Formula Produk', [
    ['1', 'Menampilkan data formula produk', 'Data formula produk berhasil tampil', 'Berhasil'],
    ['2', 'Menambah data formula produk', 'Data formula produk berhasil disimpan', 'Berhasil'],
    ['3', 'Menambah formula produk tanpa nama', 'Sistem menampilkan pesan error "Nama formula wajib diisi"', 'Berhasil'],
    ['4', 'Mengedit data formula produk', 'Data formula produk berhasil diubah', 'Berhasil']
])

# 6. Master Data: Supplier
add_table('Master Data: Supplier', [
    ['1', 'Menampilkan data supplier', 'Data supplier berhasil tampil', 'Berhasil'],
    ['2', 'Menambah data supplier', 'Data supplier berhasil disimpan', 'Berhasil'],
    ['3', 'Menambah supplier tanpa nama', 'Sistem menampilkan pesan error "Nama supplier wajib diisi"', 'Berhasil'],
    ['4', 'Mengedit data supplier', 'Data supplier berhasil diubah', 'Berhasil'],
    ['5', 'Menghapus supplier yang tidak berelasi', 'Data supplier berhasil dihapus', 'Berhasil'],
    ['6', 'Menghapus supplier yang masih berelasi dengan pembelian bahan baku', 'Sistem menolak penghapusan dan menampilkan pesan error "Data tidak dapat dihapus karena masih digunakan di transaksi lain"', 'Berhasil'],
    ['7', 'Melihat detail supplier', 'Data detail supplier berhasil tampil', 'Berhasil']
])

# 7. Master Data: User
add_table('Master Data: User', [
    ['1', 'Menampilkan data user', 'Data user berhasil tampil', 'Berhasil'],
    ['2', 'Menambah data user', 'Data user berhasil disimpan', 'Berhasil'],
    ['3', 'Menambah user tanpa username/email', 'Sistem menampilkan pesan error "Username/email wajib diisi"', 'Berhasil'],
    ['4', 'Menambah user dengan username/email yang sudah ada', 'Sistem menolak penyimpanan dan menampilkan pesan error "Username/email sudah digunakan"', 'Berhasil'],
    ['5', 'Mengedit data user', 'Data user berhasil diubah', 'Berhasil'],
    ['6', 'Menghapus data user', 'Data user berhasil dihapus', 'Berhasil'],
    ['7', 'Melihat detail user', 'Data detail user berhasil tampil', 'Berhasil']
])

# 8. Transaksi: Produk Masuk
add_table('Transaksi: Produk Masuk', [
    ['1', 'Menampilkan data produk masuk', 'Data produk masuk berhasil tampil', 'Berhasil'],
    ['2', 'Menambah data produk masuk', 'Data produk masuk berhasil disimpan', 'Berhasil'],
    ['3', 'Menambah produk masuk tanpa memilih produk', 'Sistem menampilkan pesan error "Produk wajib dipilih"', 'Berhasil'],
    ['4', 'Mengedit data produk masuk', 'Data produk masuk berhasil diubah', 'Berhasil'],
    ['5', 'Menghapus data produk masuk', 'Data produk masuk berhasil dihapus', 'Berhasil']
])

# 9. Transaksi: Produk Keluar
add_table('Transaksi: Produk Keluar', [
    ['1', 'Menampilkan data produk keluar', 'Data produk keluar berhasil tampil', 'Berhasil'],
    ['2', 'Menambah data produk keluar', 'Data produk keluar berhasil disimpan', 'Berhasil'],
    ['3', 'Menambah produk keluar tanpa memilih produk', 'Sistem menampilkan pesan error "Produk wajib dipilih"', 'Berhasil'],
    ['4', 'Menambah produk keluar dengan jumlah melebihi stok', 'Sistem menolak penyimpanan dan menampilkan pesan error "Stok tidak mencukupi"', 'Berhasil'],
    ['5', 'Mengedit data produk keluar', 'Data produk keluar berhasil diubah', 'Berhasil'],
    ['6', 'Menghapus data produk keluar', 'Data produk keluar berhasil dihapus', 'Berhasil']
])

# 10. Transaksi: Permintaan Produk
add_table('Transaksi: Permintaan Produk', [
    ['1', 'Menampilkan data permintaan produk', 'Data permintaan produk berhasil tampil', 'Berhasil'],
    ['2', 'Menambah data permintaan produk', 'Data permintaan produk berhasil disimpan', 'Berhasil'],
    ['3', 'Menambah permintaan produk tanpa memilih produk', 'Sistem menampilkan pesan error "Produk wajib dipilih"', 'Berhasil'],
    ['4', 'Mengedit data permintaan produk', 'Data permintaan produk berhasil diubah', 'Berhasil'],
    ['5', 'Menghapus data permintaan produk', 'Data permintaan produk berhasil dihapus', 'Berhasil']
])

# 11. Transaksi: Bahan Baku Masuk
add_table('Transaksi: Bahan Baku Masuk', [
    ['1', 'Menampilkan data bahan baku masuk', 'Data bahan baku masuk berhasil tampil', 'Berhasil'],
    ['2', 'Menambah data bahan baku masuk', 'Data bahan baku masuk berhasil disimpan', 'Berhasil'],
    ['3', 'Menambah bahan baku masuk tanpa memilih bahan baku', 'Sistem menampilkan pesan error "Bahan baku wajib dipilih"', 'Berhasil'],
    ['4', 'Mengedit data bahan baku masuk', 'Data bahan baku masuk berhasil diubah', 'Berhasil'],
    ['5', 'Menghapus data bahan baku masuk', 'Data bahan baku masuk berhasil dihapus', 'Berhasil']
])

# 12. Transaksi: Bahan Baku Keluar
add_table('Transaksi: Bahan Baku Keluar', [
    ['1', 'Menampilkan data bahan baku keluar', 'Data bahan baku keluar berhasil tampil', 'Berhasil'],
    ['2', 'Menambah data bahan baku keluar', 'Data bahan baku keluar berhasil disimpan', 'Berhasil'],
    ['3', 'Menambah bahan baku keluar tanpa memilih bahan baku', 'Sistem menampilkan pesan error "Bahan baku wajib dipilih"', 'Berhasil'],
    ['4', 'Menambah bahan baku keluar dengan jumlah melebihi stok', 'Sistem menolak penyimpanan dan menampilkan pesan error "Stok tidak mencukupi"', 'Berhasil'],
    ['5', 'Mengedit data bahan baku keluar', 'Data bahan baku keluar berhasil diubah', 'Berhasil'],
    ['6', 'Menghapus data bahan baku keluar', 'Data bahan baku keluar berhasil dihapus', 'Berhasil']
])

# 13. Transaksi: Pembelian Bahan Baku
add_table('Transaksi: Pembelian Bahan Baku', [
    ['1', 'Menampilkan data pembelian bahan baku', 'Data pembelian bahan baku berhasil tampil', 'Berhasil'],
    ['2', 'Menambah data pembelian bahan baku', 'Data pembelian bahan baku berhasil disimpan', 'Berhasil'],
    ['3', 'Menambah pembelian bahan baku tanpa memilih bahan baku', 'Sistem menampilkan pesan error "Bahan baku wajib dipilih"', 'Berhasil'],
    ['4', 'Mengedit data pembelian bahan baku', 'Data pembelian bahan baku berhasil diubah', 'Berhasil'],
    ['5', 'Menghapus pembelian bahan baku yang tidak berelasi', 'Data pembelian bahan baku berhasil dihapus', 'Berhasil'],
    ['6', 'Menghapus pembelian bahan baku yang masih berelasi', 'Sistem menolak penghapusan dan menampilkan pesan error "Data tidak dapat dihapus karena masih digunakan di transaksi lain"', 'Berhasil']
])

# 14. Produksi: Hasil Produksi
add_table('Produksi: Hasil Produksi', [
    ['1', 'Menampilkan data hasil produksi', 'Data hasil produksi berhasil tampil', 'Berhasil'],
    ['2', 'Menambah data hasil produksi', 'Data hasil produksi berhasil disimpan', 'Berhasil'],
    ['3', 'Menambah hasil produksi tanpa memilih produk', 'Sistem menampilkan pesan error "Produk wajib dipilih"', 'Berhasil'],
    ['4', 'Mengedit data hasil produksi', 'Data hasil produksi berhasil diubah', 'Berhasil'],
    ['5', 'Menghapus data hasil produksi', 'Data hasil produksi berhasil dihapus', 'Berhasil']
])

# 15. Produksi: Jadwal Produksi
add_table('Produksi: Jadwal Produksi', [
    ['1', 'Menampilkan data jadwal produksi', 'Data jadwal produksi berhasil tampil', 'Berhasil'],
    ['2', 'Menambah data jadwal produksi', 'Data jadwal produksi berhasil disimpan', 'Berhasil'],
    ['3', 'Menambah jadwal produksi tanpa memilih produk', 'Sistem menampilkan pesan error "Produk wajib dipilih"', 'Berhasil'],
    ['4', 'Mengedit data jadwal produksi', 'Data jadwal produksi berhasil diubah', 'Berhasil'],
    ['5', 'Menghapus jadwal produksi yang tidak berelasi', 'Data jadwal produksi berhasil dihapus', 'Berhasil'],
    ['6', 'Menghapus jadwal produksi yang masih berelasi', 'Sistem menolak penghapusan dan menampilkan pesan error "Data tidak dapat dihapus karena masih digunakan di transaksi lain"', 'Berhasil']
])

# 16. SPK (WMA + EDD)
add_table('SPK (WMA + EDD)', [
    ['1', 'Menampilkan rekomendasi jadwal produksi', 'Rencana jadwal produksi berhasil tampil', 'Berhasil'],
    ['2', 'Melakukan pencarian/filter produk', 'Data rekomendasi sesuai filter berhasil tampil', 'Berhasil'],
    ['3', 'Approve jadwal produksi', 'Jadwal produksi berhasil di-approve', 'Berhasil'],
    ['4', 'Approve tanpa memilih data', 'Sistem menampilkan pesan error "Pilih minimal satu rekomendasi untuk di-approve"', 'Berhasil'],
    ['5', 'Tidak ada data penjualan/produk keluar', 'Sistem menampilkan pesan error "Belum ada rekomendasi jadwal produksi untuk bulan depan. Data penjualan/produk keluar mungkin masih kosong."', 'Berhasil']
])

# 17. Laporan
add_table('Laporan', [
    ['1', 'Menampilkan laporan bahan baku', 'Laporan bahan baku berhasil tampil', 'Berhasil'],
    ['2', 'Menampilkan laporan produk', 'Laporan produk berhasil tampil', 'Berhasil'],
    ['3', 'Menampilkan laporan jadwal produksi', 'Laporan jadwal produksi berhasil tampil', 'Berhasil'],
    ['4', 'Menampilkan laporan SPK', 'Laporan SPK berhasil tampil', 'Berhasil'],
    ['5', 'Cetak laporan ke PDF', 'Laporan berhasil dicetak ke PDF', 'Berhasil'],
    ['6', 'Cetak laporan tanpa data', 'Sistem menampilkan pesan "Tidak ada data untuk dicetak"', 'Berhasil']
])

# 18. Notifikasi & Activity Log
add_table('Notifikasi & Activity Log', [
    ['1', 'Menampilkan notifikasi', 'Notifikasi berhasil tampil', 'Berhasil'],
    ['2', 'Menampilkan activity log', 'Data activity log berhasil tampil', 'Berhasil'],
    ['3', 'Cetak activity log ke PDF', 'Activity log berhasil dicetak ke PDF', 'Berhasil']
])

# 19. Stok Menipis
add_table('Stok Menipis', [
    ['1', 'Menampilkan data stok menipis', 'Data stok menipis berhasil tampil', 'Berhasil']
])

doc.save('tabel_pengujian_blackbox.docx') 